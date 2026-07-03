"""Конструктор проектов и задач — массовое заведение в компаниях.

POST /builder/bulk — создать пачку проектов (с вложенными задачами) и
отдельных задач сразу в НЕСКОЛЬКИХ компаниях, с общими настройками (год,
направление, доска, дедлайн по умолчанию). Переиспользует сервисы создания
задач/проектов (тот же путь, что и одиночное создание), напрямую — bulk-
операция администратора, без модерационного гейта.

Также отдаёт справочники для UI: /builder/companies, /builder/directions.
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import re
from datetime import date
from decimal import Decimal, InvalidOperation
from typing import Any, Optional
from uuid import UUID

import httpx
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.security import get_current_user, require_permission
from app.database import get_db
from app.dependencies.financials_reports import FinancialsReportsServiceDep
from app.dependencies.kpi import KpiEditorServiceDep
from app.dependencies.projects import ProjectsEditorServiceDep
from app.dependencies.tasks import TasksEditorServiceDep
from app.models.company import Company, Direction
from app.models.user import User
from app.schemas.project import ProjectCreate
from app.schemas.task import TaskCreate
from app.services import ai_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/builder", tags=["builder"])

# Допустимые статусы/приоритеты (зеркало STATUSES/PRIOS на фронте Конструктора)
_VALID_STATUS = {"new", "init", "active", "quarterly", "monthly", "ongoing"}
_VALID_PRIO = {"high", "medium", "low"}
_MAX_ROWS = 300   # верхний предел строк, отдаваемых модели за один проход


# ─── справочники для UI ────────────────────────────────────────────

@router.get("/companies")
async def builder_companies(db: AsyncSession = Depends(get_db), _u: User = Depends(get_current_user)):
    rows = (await db.execute(
        select(Company.id, Company.code, Company.name_short, Company.name_ru).order_by(Company.name_ru),
    )).all()
    return {"items": [
        {"id": str(r._mapping["id"]), "code": r._mapping["code"],
         "name": r._mapping["name_short"] or r._mapping["name_ru"]}
        for r in rows
    ]}


@router.get("/directions")
async def builder_directions(db: AsyncSession = Depends(get_db), _u: User = Depends(get_current_user)):
    rows = (await db.execute(
        select(Direction.id, Direction.code, Direction.name_ru).order_by(Direction.sort_order, Direction.name_ru),
    )).all()
    return {"items": [
        {"id": str(r._mapping["id"]), "code": r._mapping["code"], "name": r._mapping["name_ru"]}
        for r in rows
    ]}


# ─── ИИ-импорт из документов ───────────────────────────────────────

def _extract_tabular(data: bytes, filename: str) -> tuple[str, str, int]:
    """Документ → компактный текст для модели. Возвращает (text, source, rows)."""
    name = (filename or "").lower()

    # Excel
    if name.endswith((".xlsx", ".xlsm", ".xls")):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        lines: list[str] = []
        rows = 0
        for ws in wb.worksheets:
            sheet_rows = list(ws.iter_rows(values_only=True))
            if not sheet_rows:
                continue
            lines.append(f"### Лист: {ws.title}")
            for r in sheet_rows:
                # склеиваем многострочные ячейки в одну строку (внутр. \n/\t → пробел),
                # обрезаем гигантские комментарии — иначе строка-ряд рвётся и раздувает ввод
                cells = [
                    "" if c is None
                    else re.sub(r"\s+", " ", str(c).strip())[:400]
                    for c in r
                ]
                if not any(cells):
                    continue
                lines.append("\t".join(cells))
                rows += 1
                if rows >= _MAX_ROWS:
                    break
            if rows >= _MAX_ROWS:
                break
        wb.close()
        return "\n".join(lines), "excel", rows

    # CSV / TSV
    if name.endswith((".csv", ".tsv", ".txt")):
        text = data.decode("utf-8-sig", errors="replace")
        all_lines = [ln for ln in text.splitlines() if ln.strip()]
        rows = len(all_lines)
        return "\n".join(all_lines[:_MAX_ROWS]), "csv", min(rows, _MAX_ROWS)

    # PDF (текстовый слой)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(parts).strip()
        if not text:
            raise HTTPException(422, "PDF без текстового слоя (вероятно скан). OCR пока не поддержан.")
        return text[:20000], "pdf", text.count("\n") + 1

    # DOCX
    if name.endswith(".docx"):
        try:
            from docx import Document  # python-docx
        except ImportError:
            raise HTTPException(422, "Формат DOCX пока не поддержан на сервере. Используйте Excel / CSV / PDF.")
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)[:20000], "docx", len(parts)

    raise HTTPException(422, f"Неподдерживаемый формат: {filename}. Excel / CSV / PDF / DOCX.")


def _extract_json_obj(raw: str) -> dict[str, Any]:
    """Достать JSON-объект из ответа модели (срезать ```-обёртку / прочий текст)."""
    s = raw.strip()
    if "```" in s:
        m = re.search(r"```(?:json)?\s*(.+?)```", s, re.DOTALL)
        if m:
            s = m.group(1).strip()
    i, j = s.find("{"), s.rfind("}")
    if i == -1 or j == -1 or j < i:
        raise HTTPException(502, "ИИ вернул ответ без JSON. Попробуйте другой файл.")
    return json.loads(s[i:j + 1])


async def _safe_ai_json(system: str, prompt: str, max_tokens: int, retries: int = 5) -> Optional[dict]:
    """Один вызов модели → JSON-объект. None при неустранимой ошибке.

    Ретраит 429 (rate limit) и 529 (overloaded) с экспоненциальной задержкой —
    чанковый маппинг шлёт несколько запросов, важно не упасть на лимите.
    Используется для чанкового маппинга: битый блок пропускается, не валит импорт.
    """
    delay = 2.0
    for attempt in range(retries):
        try:
            raw = await ai_service.complete_once(
                system=system, prompt=prompt, max_tokens=max_tokens, temperature=0.1,
            )
            return _extract_json_obj(raw)
        except httpx.HTTPStatusError as e:
            code = e.response.status_code
            if code in (429, 529) and attempt < retries - 1:
                ra = e.response.headers.get("retry-after")
                try:
                    wait = float(ra) if ra else delay
                except (TypeError, ValueError):
                    wait = delay
                await asyncio.sleep(min(wait, 30.0))
                delay *= 2
                continue
            logger.warning("ingest: AI chunk HTTP %s", code)
            return None
        except Exception:  # noqa: BLE001
            logger.warning("ingest: AI/JSON chunk failed", exc_info=True)
            return None
    return None


def _norm_status(v: Any) -> str:
    s = str(v or "").strip().lower()
    return s if s in _VALID_STATUS else "new"


def _norm_prio(v: Any) -> str:
    s = str(v or "").strip().lower()
    return s if s in _VALID_PRIO else "medium"


def _norm_date(v: Any) -> Optional[str]:
    s = str(v or "").strip()
    return s if re.fullmatch(r"\d{4}-\d{2}-\d{2}", s) else None


def _norm_task(t: dict, dir_map: dict[str, str]) -> dict:
    return {
        "title": str(t.get("title") or "").strip()[:512],
        "status": _norm_status(t.get("status")),
        "priority": _norm_prio(t.get("priority")),
        "due_date": _norm_date(t.get("due_date")),
        "assignee_email": (str(t.get("assignee_email")).strip() or None) if t.get("assignee_email") else None,
        "direction_id": dir_map.get(str(t.get("direction") or "").strip().lower()) or "",
        "comment": str(t.get("comment") or "").strip()[:2000],
    }


class IngestOut(BaseModel):
    target: str                     # ключ дашборда-цели
    target_label: str               # человекочитаемое имя дашборда
    supported: bool                 # подключён ли к авто-созданию
    confidence: float = 0.0         # уверенность классификатора 0..1
    fields: list[dict] = []         # схема полей цели (для превью)
    projects: list[dict] = []       # заполняется для projects_tasks
    standalone_tasks: list[dict] = []
    rows: list[dict] = []           # распознанные строки для прочих целей (превью)
    rows_parsed: int = 0
    source: str = ""
    notes: str = ""


@router.post("/ingest", response_model=IngestOut)
async def ingest_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("tasks.edit")),
):
    """Файл (Excel/CSV/PDF/DOCX) → ИИ-классификация дашборда + маппинг полей → превью.

    Агент СНАЧАЛА определяет, к какому дашборду относится документ (по реестру
    ingest_registry), ЗНАЯ структуру каждого дашборда, затем маппит данные в его
    поля. Ничего НЕ создаёт: возвращает структуру на превью. Для projects_tasks
    подтверждение идёт через /builder/bulk; прочие цели — пока только распознавание.
    """
    from app.services import ingest_registry as reg

    if not ai_service.is_enabled():
        raise HTTPException(503, "ИИ-импорт недоступен: ИИ-движок не настроен.")

    data = await file.read()
    if not data:
        raise HTTPException(422, "Пустой файл.")
    if len(data) > 8 * 1024 * 1024:
        raise HTTPException(413, "Файл больше 8 МБ. Разбейте на части.")

    table_text, source, rows = _extract_tabular(data, file.filename or "")
    if not table_text.strip():
        raise HTTPException(422, "Не удалось извлечь данные из файла.")

    # справочник направлений → модель оперирует именами, сервер резолвит в id
    dir_rows = (await db.execute(select(Direction.id, Direction.name_ru))).all()
    dir_map = {r._mapping["name_ru"].strip().lower(): str(r._mapping["id"]) for r in dir_rows}
    dir_names = [r._mapping["name_ru"] for r in dir_rows]

    lines = [ln for ln in table_text.split("\n") if ln.strip()]
    header_ctx = "\n".join(lines[:6])   # шапка/столбцы как контекст для каждого чанка

    # ── ШАГ 1: КЛАССИФИКАЦИЯ дашборда по сэмплу (один вызов) ──
    classify_system = (
        "Ты — агент структурирования данных платформы UzAssets (портфель госпредприятий). "
        "Определи, к какому ДАШБОРДУ относится документ. Реестр:\n"
        f"{reg.target_catalog_for_prompt()}\n\n"
        "Верни ТОЛЬКО JSON (без markdown): {\"target\": \"<key|other>\", \"confidence\": <0..1>, "
        "\"notes\": \"1-2 предложения\", "
        "\"inferred_fields\": [{\"name\": str, \"type\": \"str|number|date|enum\", \"desc\": str}]}. "
        "inferred_fields заполняй ТОЛЬКО если target=\"other\" (структуры нет в реестре)."
    )
    cls = await _safe_ai_json(classify_system, f"Документ ({source}):\n" + "\n".join(lines[:50]), 1500)
    if cls is None:
        raise HTTPException(502, "ИИ не смог классифицировать документ. Попробуйте другой файл.")

    target_key = str(cls.get("target") or "projects_tasks").strip()
    tgt = reg.BY_KEY.get(target_key)

    # ── ШАГ 2: МАППИНГ строк ЧАНКАМИ (параллельно) — устойчиво к большим файлам ──
    CHUNK = 40
    body = lines if len(lines) <= CHUNK else lines
    chunks = [body[i:i + CHUNK] for i in range(0, len(body), CHUNK)] or [body]

    if tgt is not None and tgt.key == "projects_tasks":
        map_system = (
            "Разбери строки документа в проекты и задачи UzAssets. Верни ТОЛЬКО JSON: "
            "{\"projects\":[{\"title\",\"status\",\"priority\",\"due_date\",\"direction\",\"comment\","
            "\"tasks\":[{\"title\",\"status\",\"priority\",\"due_date\",\"direction\",\"comment\"}]}],"
            "\"standalone_tasks\":[{\"title\",\"status\",\"priority\",\"due_date\",\"direction\",\"comment\"}]}.\n"
            f"status ∈ {reg.TASK_STATUS} (деф new), priority ∈ {reg.PRIORITY} (деф medium). "
            "Иерархия (пункт 1 и подпункты 1.1/1.2 …) → подзадачи в tasks; иначе standalone_tasks. "
            "Даты → YYYY-MM-DD или null. "
            f"direction ТОЛЬКО точное имя из списка или \"\": {dir_names}. "
            "comment — перенеси сюда текст из колонок «Комментарий»/«Примечание»/«Заметки»/«Результаты» "
            "(≤1500 символов), если их нет — пустая строка. "
            "Заголовки и пустые строки игнорируй. title ≤ 300 символов (только название, без комментариев). "
            "Ничего не выдумывай."
        )
        key_a, key_b = "projects", "standalone_tasks"
    else:
        if tgt is not None:
            fields_desc = ", ".join(
                f"{f.name}(" + ("|".join(f.enum) if f.enum else f.type) + ")" for f in tgt.fields
            )
            label = tgt.label
        else:
            inf_names = [str(f.get("name")).strip() for f in (cls.get("inferred_fields") or []) if str(f.get("name") or "").strip()]
            fields_desc = ", ".join(inf_names) if inf_names else "определи столбцы сам"
            label = "новая структура"
        map_system = (
            f"Разбери строки документа в JSON для дашборда «{label}». Верни ТОЛЬКО JSON: "
            "{\"rows\":[{<field>:<value>}]}. "
            f"Ключи объектов — РОВНО эти поля: {fields_desc}. "
            "enum приводи к допустимым; даты → YYYY-MM-DD или null; числа — числом. "
            "Заголовки и пустые строки игнорируй. Ничего не выдумывай."
        )
        key_a, key_b = "rows", None

    sem = asyncio.Semaphore(4)   # ограничиваем параллелизм — иначе Anthropic 429

    async def _map_chunk(chunk_lines: list[str]) -> Optional[dict]:
        prompt = f"Столбцы/шапка:\n{header_ctx}\n\nСтроки:\n" + "\n".join(chunk_lines)
        async with sem:
            return await _safe_ai_json(map_system, prompt, 8000)

    frags = await asyncio.gather(*[_map_chunk(c) for c in chunks])

    obj: dict[str, Any] = {
        "target": target_key,
        "confidence": cls.get("confidence"),
        "notes": cls.get("notes") or "",
        "inferred_fields": cls.get("inferred_fields") or [],
        "projects": [],
        "standalone_tasks": [],
        "rows": [],
    }
    ok_chunks = 0
    for frag in frags:
        if not isinstance(frag, dict):
            continue
        ok_chunks += 1
        obj[key_a].extend(frag.get(key_a) or [])
        if key_b:
            obj[key_b].extend(frag.get(key_b) or [])
    if ok_chunks == 0:
        raise HTTPException(502, "ИИ не смог распарсить данные документа (структура слишком сложная).")
    if ok_chunks < len(chunks):
        obj["notes"] = (str(obj.get("notes") or "") +
                        f" ⚠ Распознано {ok_chunks}/{len(chunks)} блоков — часть строк могла не попасть.").strip()

    target_key = str(obj.get("target") or "projects_tasks").strip()
    tgt = reg.BY_KEY.get(target_key)   # None → неизвестная/новая структура (target=other)
    try:
        confidence = max(0.0, min(1.0, float(obj.get("confidence") or 0)))
    except (TypeError, ValueError):
        confidence = 0.0

    projects: list[dict] = []
    standalone: list[dict] = []
    out_rows: list[dict] = []

    if tgt is not None and tgt.key == "projects_tasks":
        for p in obj.get("projects") or []:
            if not str(p.get("title") or "").strip():
                continue
            projects.append({
                "title": str(p.get("title")).strip()[:512],
                "status": _norm_status(p.get("status")),
                "priority": _norm_prio(p.get("priority")),
                "due_date": _norm_date(p.get("due_date")),
                "direction_id": dir_map.get(str(p.get("direction") or "").strip().lower()) or "",
                "comment": str(p.get("comment") or "").strip()[:2000],
                "tasks": [_norm_task(t, dir_map) for t in (p.get("tasks") or []) if str(t.get("title") or "").strip()],
            })
        standalone = [_norm_task(t, dir_map) for t in (obj.get("standalone_tasks") or []) if str(t.get("title") or "").strip()]
        if not projects and not standalone:
            raise HTTPException(422, "ИИ не нашёл проектов/задач в документе.")
        out_key, out_label, out_supported = "projects_tasks", tgt.label, True
        out_fields = reg.fields_meta("projects_tasks")
    else:
        # известный дашборд (не projects_tasks) ИЛИ новая структура (other) —
        # агент адаптируется: схема берётся из реестра либо из inferred_fields.
        if tgt is not None:
            out_key, out_label, out_supported = tgt.key, tgt.label, tgt.supported
            out_fields = reg.fields_meta(tgt.key)
        else:
            inferred = [
                {"name": str(f.get("name")).strip(), "type": str(f.get("type") or "str"),
                 "desc": str(f.get("desc") or ""), "enum": []}
                for f in (obj.get("inferred_fields") or [])
                if str(f.get("name") or "").strip()
            ]
            out_key, out_label, out_supported = "other", "Новая структура (распознана автоматически)", False
            out_fields = inferred
        allowed = {f["name"] for f in out_fields}
        for r in obj.get("rows") or []:
            if not isinstance(r, dict):
                continue
            # для известного дашборда — фильтруем по схеме; для other (allowed может быть
            # пустым, если поля не выведены) — принимаем все ключи строки.
            clean = {
                str(k): ("" if v is None else str(v).strip())
                for k, v in r.items()
                if (not allowed) or k in allowed
            }
            if any(clean.values()):
                out_rows.append(clean)
        # other без выведённых полей — берём ключи из первой строки
        if out_key == "other" and not out_fields and out_rows:
            out_fields = [{"name": k, "type": "str", "desc": "", "enum": []} for k in out_rows[0].keys()]
        if not out_rows:
            raise HTTPException(422, f"ИИ отнёс документ к «{out_label}», но не распознал строки.")

    return IngestOut(
        target=out_key,
        target_label=out_label,
        supported=out_supported,
        confidence=confidence,
        fields=out_fields,
        projects=projects,
        standalone_tasks=standalone,
        rows=out_rows,
        rows_parsed=rows,
        source=source,
        notes=str(obj.get("notes") or "")[:400],
    )


# ─── bulk-схемы ────────────────────────────────────────────────────

class BulkTask(BaseModel):
    title: str = Field(..., min_length=1, max_length=512)
    status: str = "new"
    priority: str = "medium"
    due_date: Optional[date] = None
    assignee_email: Optional[str] = None
    direction_id: Optional[UUID] = None
    comment: Optional[str] = None       # переносится из колонки «Комментарий» при импорте


class BulkProject(BaseModel):
    title: str = Field(..., min_length=1, max_length=512)
    status: str = "new"
    priority: str = "medium"
    due_date: Optional[date] = None
    direction_id: Optional[UUID] = None
    comment: Optional[str] = None
    tasks: list[BulkTask] = Field(default_factory=list)


class BulkCommon(BaseModel):
    portfolio_year: Optional[int] = None
    direction_id: Optional[UUID] = None       # направление по умолчанию
    board_id: Optional[UUID] = None
    due_date: Optional[date] = None           # дедлайн по умолчанию


class BulkRequest(BaseModel):
    company_ids: list[UUID] = Field(default_factory=list)
    common: BulkCommon = Field(default_factory=BulkCommon)
    projects: list[BulkProject] = Field(default_factory=list)
    standalone_tasks: list[BulkTask] = Field(default_factory=list)


def _pick(*vals):
    for v in vals:
        if v is not None:
            return v
    return None


async def _enforce_company_scope(db: AsyncSession, user, cids) -> None:
    """Security (audit H-2): запретить bulk-запись в компании ВНЕ scope актора.
    Раньше /builder/bulk* писали в любую company_id, проверяя только глобальное
    право (tasks.edit/kpi.edit/financials.edit). Теперь сверяем целевые компании
    с allowed_company_ids. cids — итерируемое company_id (UUID/str/None); None
    (без привязки к компании) пропускаем. scope==None → owner/companies.view_all."""
    from app.core.access import allowed_company_ids
    scope = await allowed_company_ids(db, user)
    if scope is None:
        return
    allowed = {str(x) for x in scope}
    for cid in cids:
        if cid is not None and str(cid) not in allowed:
            raise HTTPException(403, "Нет доступа к одной из выбранных компаний.")


@router.post("/bulk")
async def bulk_create(
    body: BulkRequest,
    tasks_svc: TasksEditorServiceDep,
    projects_svc: ProjectsEditorServiceDep,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("tasks.edit")),
):
    """Массовое создание проектов+задач в выбранных компаниях."""
    c = body.common
    targets = body.company_ids or [None]   # если не выбрано — без привязки к компании
    await _enforce_company_scope(db, user, targets)   # H-2: scope-guard
    proj_n = 0
    task_n = 0
    # (kind, parent_id, body) — комментарии из импорта, создаём после сущностей
    pending_comments: list[tuple[str, UUID, str]] = []

    def _stash_comment(kind: str, pid: UUID, text: Optional[str]) -> None:
        body_txt = (text or "").strip()
        if body_txt:
            pending_comments.append((kind, pid, body_txt[:5000]))

    for cid in targets:
        # проекты (+ вложенные задачи)
        for p in body.projects:
            pc = ProjectCreate(
                title=p.title, status=p.status, priority=p.priority,
                company_id=cid, portfolio_year=c.portfolio_year, board_id=c.board_id,
                direction_id=_pick(p.direction_id, c.direction_id),
                due_date=_pick(p.due_date, c.due_date),
            )
            detail, _info = await projects_svc.create_project(pc, creator_id=user.id)
            proj_n += 1
            pid = UUID(str(detail.id)) if not isinstance(detail.id, UUID) else detail.id
            _stash_comment("project", pid, p.comment)
            for t in p.tasks:
                tc = TaskCreate(
                    title=t.title, status=t.status, priority=t.priority,
                    company_id=cid, project_id=pid, portfolio_year=c.portfolio_year,
                    board_id=c.board_id,
                    direction_id=_pick(t.direction_id, p.direction_id, c.direction_id),
                    due_date=_pick(t.due_date, c.due_date),
                    assignee_email=t.assignee_email,
                )
                created, _ = await tasks_svc.create_task(tc, creator_id=user.id)
                task_n += 1
                _stash_comment("task", created.id, t.comment)

        # отдельные задачи (без проекта)
        for t in body.standalone_tasks:
            tc = TaskCreate(
                title=t.title, status=t.status, priority=t.priority,
                company_id=cid, portfolio_year=c.portfolio_year, board_id=c.board_id,
                direction_id=_pick(t.direction_id, c.direction_id),
                due_date=_pick(t.due_date, c.due_date),
                assignee_email=t.assignee_email,
            )
            created, _ = await tasks_svc.create_task(tc, creator_id=user.id)
            task_n += 1
            _stash_comment("task", created.id, t.comment)

    # комментарии из импорта — прямые вставки (без модерации/нотификаций), один commit
    comment_n = 0
    if pending_comments:
        from app.models.project import ProjectComment
        from app.models.task import TaskComment
        for kind, pid, body_txt in pending_comments:
            if kind == "task":
                db.add(TaskComment(task_id=pid, author_id=user.id, body=body_txt))
            else:
                db.add(ProjectComment(project_id=pid, author_id=user.id, body=body_txt))
            comment_n += 1
        await db.commit()

    return {
        "companies": len([t for t in targets if t is not None]) or 1,
        "projects_created": proj_n,
        "tasks_created": task_n,
        "comments_created": comment_n,
    }


# ─── KPI bulk (ИИ-импорт) ──────────────────────────────────────────

class BulkKpiRow(BaseModel):
    company: str = ""
    indicator: str = ""
    unit: Optional[str] = None
    weight: Optional[str] = None
    plan: Optional[str] = None
    fact: Optional[str] = None
    period: Optional[str] = None


class BulkKpiRequest(BaseModel):
    year: int
    manager_title: str = "Импорт KPI"
    rows: list[BulkKpiRow] = Field(default_factory=list)


def _norm_co(s: Any) -> str:
    return re.sub(r"\s+", " ", str(s or "").strip().lower())


@router.post("/bulk-kpi")
async def bulk_create_kpi(
    body: BulkKpiRequest,
    kpi_svc: KpiEditorServiceDep,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("kpi.edit")),
):
    """Массовое заведение KPI-индикаторов из распознанных строк.

    Каждая строка содержит имя компании → резолвим в company_id, группируем по
    компании и аддитивно добавляем индикаторы под менеджера manager_title за year.
    """
    co_rows = (await db.execute(
        select(Company.id, Company.code, Company.name_short, Company.name_ru),
    )).all()
    exact: dict[str, str] = {}
    co_list: list[tuple[str, str]] = []
    for r in co_rows:
        cid = str(r._mapping["id"])
        label = r._mapping["name_short"] or r._mapping["name_ru"] or ""
        co_list.append((cid, label))
        for key in (r._mapping["code"], r._mapping["name_short"], r._mapping["name_ru"]):
            if key:
                exact[_norm_co(key)] = cid

    def resolve(name: str) -> Optional[str]:
        n = _norm_co(name)
        if not n:
            return None
        if n in exact:
            return exact[n]
        for cid, label in co_list:              # подстрочный фолбэк
            ln = _norm_co(label)
            if ln and (n in ln or ln in n):
                return cid
        return None

    grouped: dict[str, list[dict]] = {}
    unresolved: list[str] = []
    for row in body.rows:
        if not str(row.indicator or "").strip():
            continue
        cid = resolve(row.company)
        if cid is None:
            if str(row.company or "").strip():
                unresolved.append(row.company)
            continue
        grouped.setdefault(cid, []).append({
            "name": row.indicator, "unit": row.unit,
            "weight": row.weight, "plan": row.plan, "fact": row.fact,
        })

    if not grouped:
        raise HTTPException(
            422,
            "Не удалось сопоставить ни одну компанию из документа. "
            + (f"Не распознаны: {', '.join(sorted(set(unresolved))[:8])}" if unresolved else ""),
        )

    await _enforce_company_scope(db, user, grouped.keys())   # H-2: scope-guard

    total_ind = 0
    for cid, inds in grouped.items():
        res = await kpi_svc.bulk_add_indicators(UUID(cid), body.year, body.manager_title, inds)
        total_ind += res["indicators_added"]

    return {
        "companies": len(grouped),
        "indicators_created": total_ind,
        "unresolved": sorted(set(unresolved)),
    }


# ─── Финансы bulk (ИИ-импорт) ──────────────────────────────────────

class BulkFinRow(BaseModel):
    company: str = ""
    article: str = ""
    value: Optional[str] = None
    report_type: Optional[str] = None
    standard: Optional[str] = None
    year: Optional[str] = None
    currency: Optional[str] = None


class BulkFinRequest(BaseModel):
    default_year: int
    default_standard: str = "IFRS"
    default_report_type: str = "PL"
    default_currency: str = "UZS"
    rows: list[BulkFinRow] = Field(default_factory=list)


async def _company_resolver(db: AsyncSession):
    """Замыкание-резолвер имени компании → company_id (точный + подстрочный)."""
    co_rows = (await db.execute(
        select(Company.id, Company.code, Company.name_short, Company.name_ru),
    )).all()
    exact: dict[str, str] = {}
    co_list: list[tuple[str, str]] = []
    for r in co_rows:
        cid = str(r._mapping["id"])
        label = r._mapping["name_short"] or r._mapping["name_ru"] or ""
        co_list.append((cid, label))
        for key in (r._mapping["code"], r._mapping["name_short"], r._mapping["name_ru"]):
            if key:
                exact[_norm_co(key)] = cid

    def resolve(name: str) -> Optional[str]:
        n = _norm_co(name)
        if not n:
            return None
        if n in exact:
            return exact[n]
        for cid, label in co_list:
            ln = _norm_co(label)
            if ln and (n in ln or ln in n):
                return cid
        return None
    return resolve


def _to_decimal(v: Any) -> Optional[Decimal]:
    s = str(v or "").strip().replace(" ", "").replace(" ", "").replace(",", ".")
    if not s:
        return None
    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return None


@router.post("/bulk-financials")
async def bulk_create_financials(
    body: BulkFinRequest,
    fin_svc: FinancialsReportsServiceDep,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(require_permission("financials.edit")),
):
    """Массовое заведение строк финотчётов из распознанных строк.

    Имя компании в каждой строке → company_id. report_type/standard/year/currency
    берём из строки, иначе из дефолтов запроса. Группировка по отчёту — в сервисе
    (get-or-create отчёт + аддитивные строки).
    """
    resolve = await _company_resolver(db)
    std_ok = {"IFRS", "NSBU"}
    rt_ok = {"PL", "BS", "CF"}

    rows: list[dict] = []
    unresolved: list[str] = []
    for row in body.rows:
        if not str(row.article or "").strip():
            continue
        cid = resolve(row.company)
        if cid is None:
            if str(row.company or "").strip():
                unresolved.append(row.company)
            continue
        std = str(row.standard or "").strip().upper() or body.default_standard
        std = std if std in std_ok else body.default_standard
        rt = str(row.report_type or "").strip().upper() or body.default_report_type
        rt = rt if rt in rt_ok else body.default_report_type
        try:
            yr = int(str(row.year or "").strip() or body.default_year)
        except (TypeError, ValueError):
            yr = body.default_year
        rows.append({
            "company_id": UUID(cid),
            "year": yr,
            "quarter": None,
            "standard": std,
            "report_type": rt,
            "currency": str(row.currency or "").strip().upper() or body.default_currency,
            # канон платформы: value в МЛРД сум → unit_scale=1e9 (аудит P1)
            "unit_scale": 1_000_000_000,
            "article": row.article,
            "value": _to_decimal(row.value),
        })

    if not rows:
        raise HTTPException(
            422,
            "Не удалось сопоставить ни одну компанию из документа. "
            + (f"Не распознаны: {', '.join(sorted(set(unresolved))[:8])}" if unresolved else ""),
        )

    await _enforce_company_scope(db, user, {r["company_id"] for r in rows})   # H-2: scope-guard

    res = await fin_svc.bulk_add_lines(rows, db, user)
    return {
        "reports": res["reports"],
        "lines_created": res["lines_added"],
        "unresolved": sorted(set(unresolved)),
    }
