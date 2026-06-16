"""Knowledge base (RAG) — загрузка/список/удаление документов.

Текст извлекается, режется на чанки и индексируется Postgres FTS (tsv заполняет
триггер). Ассистент ищет по базе инструментом search_knowledge_base.
Управление — только super-admin/owner.
"""
from __future__ import annotations

import io
import logging
import re
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user, get_db
from app.core.security import is_super_admin
from app.models.knowledge import KnowledgeChunk, KnowledgeDoc
from app.models.user import User
from app.services import embeddings

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/knowledge", tags=["knowledge"])

MAX_BYTES = 8_000_000
CHUNK_SIZE = 900


def require_admin(user: User = Depends(get_current_user)) -> User:
    if not is_super_admin(user):
        raise HTTPException(status.HTTP_403_FORBIDDEN, "Доступ только для администратора")
    return user


def _extract_text(filename: Optional[str], raw: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            return "\n\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            return ""
    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""
    if name.endswith((".xlsx", ".xlsm")):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts: list[str] = []
            for ws in wb.worksheets:
                parts.append(f"# Лист: {ws.title}")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""
    for enc in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _chunk(text: str, size: int = CHUNK_SIZE) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    chunks: list[str] = []
    buf = ""
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        if len(buf) + len(para) + 2 <= size:
            buf = (buf + "\n\n" + para).strip()
        else:
            if buf:
                chunks.append(buf)
                buf = ""
            if len(para) <= size:
                buf = para
            else:
                for i in range(0, len(para), size):
                    chunks.append(para[i:i + size])
    if buf:
        chunks.append(buf)
    return chunks[:2000]


async def _embed_chunks(db: AsyncSession, chunk_ids: list, texts: list[str]) -> int:
    """Считает эмбеддинги Voyage и пишет их в knowledge_chunk.embedding.

    Изолировано в SAVEPOINT: при сбое (нет ключа / провайдер недоступен /
    нет pgvector) откатывается только запись векторов — чанки и FTS остаются.
    Возвращает число проиндексированных чанков (0, если слой отключён).
    """
    if not embeddings.is_enabled() or not texts:
        return 0
    try:
        vecs = await embeddings.embed_documents(texts)
        async with db.begin_nested():
            for cid, vec in zip(chunk_ids, vecs):
                await db.execute(
                    text("UPDATE knowledge_chunk SET embedding = CAST(:v AS vector) "
                         "WHERE id = :id"),
                    {"v": embeddings.to_pgvector(vec), "id": cid},
                )
        return len(vecs)
    except Exception as e:  # noqa: BLE001
        logger.warning("knowledge: эмбеддинги не записаны (FTS работает): %s", e)
        return 0


@router.post("/upload")
async def upload(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    raw = await file.read()
    if len(raw) > MAX_BYTES:
        raise HTTPException(status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, "Файл больше 8 МБ")
    body = _extract_text(file.filename, raw)
    chunks = _chunk(body)
    if not chunks:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "Не удалось извлечь текст из файла")
    doc = KnowledgeDoc(
        title=(title or file.filename or "Документ")[:512],
        filename=file.filename,
        content_type=file.content_type,
        char_count=len(body),
        chunk_count=len(chunks),
        uploaded_by=user.id,
    )
    db.add(doc)
    await db.flush()
    objs = [KnowledgeChunk(doc_id=doc.id, chunk_index=i, content=c) for i, c in enumerate(chunks)]
    for o in objs:
        db.add(o)
    await db.flush()  # назначить id чанкам до записи эмбеддингов
    embedded = await _embed_chunks(db, [o.id for o in objs], chunks)
    await db.commit()
    return {
        "id": str(doc.id), "title": doc.title, "chunks": len(chunks),
        "chars": len(body), "embedded": embedded,
    }


@router.get("")
async def list_docs(
    user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    rows = (await db.execute(
        select(KnowledgeDoc).order_by(KnowledgeDoc.created_at.desc()),
    )).scalars().all()
    return [{
        "id": str(d.id), "title": d.title, "filename": d.filename,
        "chunks": d.chunk_count, "chars": d.char_count,
        "created_at": d.created_at.isoformat() if d.created_at else None,
    } for d in rows]


@router.post("/reindex")
async def reindex(
    full: bool = False,
    user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    """Бэкфилл семантических эмбеддингов для уже загруженных чанков.

    По умолчанию заполняет только пустые (embedding IS NULL). full=true —
    пересчитывает все (например, после смены модели). Требует VOYAGE_API_KEY и
    pgvector; иначе вернёт ok=false с пояснением.
    """
    if not embeddings.is_enabled():
        return {"ok": False, "reason": "VOYAGE_API_KEY не задан — семантический слой отключён"}
    try:
        cursor = None
        processed = 0
        while True:
            cond, params = [], {"lim": 96}
            if not full:
                cond.append("embedding IS NULL")
            if cursor is not None:
                cond.append("id > :cursor")
                params["cursor"] = cursor
            where = ("WHERE " + " AND ".join(cond)) if cond else ""
            rows = (await db.execute(text(
                f"SELECT id, content FROM knowledge_chunk {where} ORDER BY id LIMIT :lim",
            ), params)).all()
            if not rows:
                break
            ids = [r[0] for r in rows]
            texts = [r[1] or "" for r in rows]
            n = await _embed_chunks(db, ids, texts)
            await db.commit()
            processed += n
            cursor = ids[-1]
            if n == 0:  # сбой эмбеддинга — не зацикливаемся
                break
        return {"ok": True, "embedded": processed}
    except Exception as e:  # noqa: BLE001
        logger.warning("knowledge: reindex не выполнен: %s", e)
        return {"ok": False, "reason": f"Семантический слой недоступен: {e}"}


@router.delete("/{doc_id}")
async def delete_doc(
    doc_id: UUID,
    user: User = Depends(require_admin),
    db: AsyncSession = Depends(get_db),
):
    d = await db.get(KnowledgeDoc, doc_id)
    if d:
        await db.delete(d)
        await db.commit()
    return {"ok": True}
